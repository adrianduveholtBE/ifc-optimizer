#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""BIM Engines instruktionsmall — docx + txt ur en spec-JSON.

Pythonport av skillen bim-engine-instructions (samma designspraak: mintgron
accent #85B09A, nastan svart #1A1A1A, gra #6B6B6B, Arial, A4, en sida).
Anvands eftersom node inte finns installerat pa den har maskinen.

  python tools/build_instructions.py spec.json utdata_basnamn [logo.png]
"""
import io, json, os, re, sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MINT = RGBColor(0x85, 0xB0, 0x9A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x6B, 0x6B, 0x6B)


def border(par, edge, size, color, space=2):
    p = par._p.get_or_add_pPr()
    bd = p.find(qn('w:pBdr'))
    if bd is None:
        bd = OxmlElement('w:pBdr')
        p.append(bd)
    e = OxmlElement('w:' + edge)
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), str(size))
    e.set(qn('w:space'), str(space))
    e.set(qn('w:color'), color)
    bd.append(e)


def spacing(par, before=None, after=None):
    pf = par.paragraph_format
    if before is not None:
        pf.space_before = Pt(before / 20.0)
    if after is not None:
        pf.space_after = Pt(after / 20.0)
    pf.line_spacing = 1.10


def runs(par, text, size=9.5, color=BLACK, bold=False, italic=False):
    for part in re.split(r'(\*\*[^*]+\*\*)', str(text)):
        if not part:
            continue
        b = part.startswith('**') and part.endswith('**')
        r = par.add_run(part[2:-2] if b else part)
        r.font.name = 'Arial'
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.bold = bool(b or bold)
        r.italic = italic


def strip_bold(s):
    return str(s).replace('**', '')


def wrap(s, width=58, indent=''):
    words = strip_bold(s).split()
    lines, line = [], ''
    for w in words:
        if len((line + ' ' + w).strip()) > width:
            lines.append(line)
            line = w
        else:
            line = (line + ' ' + w).strip()
    if line:
        lines.append(line)
    return '\n'.join((l if i == 0 else indent + l) for i, l in enumerate(lines))


def build_docx(spec, out, logo):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Arial'
    st.font.size = Pt(9.5)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for a in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, a, Cm(1.7))

    if logo and os.path.exists(logo):
        p = doc.add_paragraph()
        spacing(p, after=40)
        p.add_run().add_picture(logo, width=Cm(1.5), height=Cm(1.5))

    p = doc.add_paragraph()
    spacing(p, after=0)
    runs(p, spec.get('app_name', 'App'), size=20, bold=True)

    if spec.get('tagline'):
        p = doc.add_paragraph()
        spacing(p, after=60)
        runs(p, spec['tagline'], size=9.5, color=GREY)

    p = doc.add_paragraph()
    spacing(p, after=160)
    runs(p, 'BIM ENGINE AB', size=8, color=MINT, bold=True)
    border(p, 'bottom', 12, '1A1A1A', 1)

    for s in spec.get('sections', []):
        p = doc.add_paragraph()
        spacing(p, before=200, after=90)
        runs(p, s.get('heading', ''), size=12, bold=True)
        border(p, 'bottom', 6, '85B09A', 2)

        kind = s.get('type', 'paragraph')
        if kind == 'paragraph':
            for para in s.get('body', []):
                q = doc.add_paragraph()
                spacing(q, after=80)
                runs(q, para)
        elif kind == 'steps':
            items = s.get('items', [])
            for i, it in enumerate(items):
                q = doc.add_paragraph()
                spacing(q, after=100 if i == len(items) - 1 else 55)
                q.paragraph_format.left_indent = Pt(26)
                q.paragraph_format.first_line_indent = Pt(-15)
                runs(q, '%d.  ' % (i + 1), bold=True)
                runs(q, it)
        elif kind == 'bullets':
            items = s.get('items', [])
            for i, it in enumerate(items):
                q = doc.add_paragraph()
                spacing(q, after=100 if i == len(items) - 1 else 55)
                q.paragraph_format.left_indent = Pt(21)
                q.paragraph_format.first_line_indent = Pt(-12)
                runs(q, '•  ', color=MINT, bold=True)
                runs(q, it)

    if spec.get('contact'):
        p = doc.add_paragraph()
        spacing(p, before=200)
        border(p, 'top', 6, 'CCCCCC', 4)
        runs(p, spec['contact'] + '  ·  BIM Engine AB', size=8, color=GREY, italic=True)

    doc.save(out + '.docx')
    print('skrev %s.docx (%d B)' % (out, os.path.getsize(out + '.docx')))


def build_txt(spec, out):
    bar = '=' * 56
    t = [bar, '  ' + spec.get('app_name', 'App').upper() + '  —  BIM ENGINE AB']
    if spec.get('tagline'):
        t.append('  ' + wrap(spec['tagline'], 52, '  '))
    t.append(bar)
    t.append('')
    t.append('')
    for s in spec.get('sections', []):
        h = s.get('heading', '').upper()
        t.append(h)
        t.append('-' * len(h))
        kind = s.get('type', 'paragraph')
        if kind == 'paragraph':
            for para in s.get('body', []):
                t.append(wrap(para))
                t.append('')
        elif kind == 'steps':
            for i, it in enumerate(s.get('items', [])):
                t.append('%d. %s' % (i + 1, wrap(it, 56, '   ')))
                t.append('')
        elif kind == 'bullets':
            for it in s.get('items', []):
                t.append('- ' + wrap(it, 56, '  '))
                t.append('')
        t.append('')
    if spec.get('contact'):
        t.append('-' * 56)
        t.append(strip_bold(spec['contact']))
        t.append('BIM Engine AB')
        t.append('-' * 56)
    io.open(out + '.txt', 'w', encoding='utf-8', newline='\r\n').write('\n'.join(t) + '\n')
    print('skrev %s.txt (%d B)' % (out, os.path.getsize(out + '.txt')))


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        return 1
    spec = json.load(io.open(sys.argv[1], encoding='utf-8'))
    out = sys.argv[2]
    logo = sys.argv[3] if len(sys.argv) > 3 else None
    build_docx(spec, out, logo)
    build_txt(spec, out)
    return 0


if __name__ == '__main__':
    sys.exit(main())
